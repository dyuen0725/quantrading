from sklearn.feature_extraction.text import CountVectorizer
import pandas as pd
from docx import Document

#setting up the corpus
doc= Document('HK_article_1.docx')
doc1=""
for para in doc.paragraphs:
	doc1 += para.text
doc1 = doc1.replace('\xa0','')    

doc= Document('HK_article_2.docx')
doc2=""
for para in doc.paragraphs:
	doc2 += para.text
doc2 = doc2.replace('\xa0','')    

doc= Document('HK_article_3.docx')  
doc3=""
for para in doc.paragraphs:
	doc3 += para.text
doc3 = doc3.replace('\xa0','')   

doc= Document('HK_article_4.docx')
doc4=""
for para in doc.paragraphs:
	doc4 += para.text
doc4 = doc4.replace('\xa0','')   

c = [doc1, doc2, doc3, doc4]

#creating the BoW
v = CountVectorizer(stop_words='english')
bow = v.fit_transform(c)
df = pd.DataFrame(bow.A)
df.columns = v.get_feature_names()

#BoW result analysis
count = df.sum().sort_values(ascending = False)
print(count.head(15))